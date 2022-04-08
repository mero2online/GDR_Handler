from HelperFunc import resource_path, readLocalFile
from docx import Document
from docx.shared import RGBColor
from docx.shared import Pt


def HandleWord(allText):
    document = Document(resource_path(
        'input\\GEOLOGICAL_DESCRIPTION_REPORT.docx'))

    for idx, x in enumerate(allText):
        paragraph = document.add_paragraph()
        paragraph.alignment = 2  # for left, 1 for center, 2 right, 3 justify

        run = paragraph.add_run(x[0]+'\n\n\n')
        font = run.font
        font.color.rgb = RGBColor(255, 0, 0)
        font.name = 'Times New Roman'
        font.size = Pt(14)
        run.bold = True
        run.underline = True
        for i in x[1]:
            rock = paragraph.add_run(i+'\n\n\n')
            font = rock.font
            font.color.rgb = RGBColor(255, 0, 0)
            font.name = 'Times New Roman'
            font.size = Pt(11)
            rock.bold = True

    document.save(resource_path('output\\GEOLOGICAL_DESCRIPTION_REPORT.docx'))
